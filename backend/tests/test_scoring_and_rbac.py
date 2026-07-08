import io

import pytest

JD = "Python engineer: FastAPI, PostgreSQL, Docker, AWS, NLP, Machine Learning."

STRONG = (
    b"Strong Match\nstrong@example.com\n8 years experience.\n"
    b"Skills: Python, FastAPI, PostgreSQL, Docker, AWS, NLP, Machine Learning\n"
)
WEAK = (
    b"Weak Match\nweak@example.com\n2 years experience.\n"
    b"Skills: Java, React, MySQL\n"
)


def _setup_job_with_two_candidates(client, headers) -> str:
    job_id = client.post(
        "/v1/jobs", json={"title": "Backend", "description": JD}, headers=headers
    ).json()["id"]
    files = [
        ("files", ("strong.txt", io.BytesIO(STRONG), "text/plain")),
        ("files", ("weak.txt", io.BytesIO(WEAK), "text/plain")),
    ]
    client.post(f"/v1/jobs/{job_id}/resumes/upload", files=files, headers=headers)
    return job_id


def test_ranked_candidates_sorted_by_final_score(client, recruiter_headers):
    job_id = _setup_job_with_two_candidates(client, recruiter_headers)
    ranked = client.get(
        f"/v1/jobs/{job_id}/ranked-candidates", headers=recruiter_headers
    ).json()

    assert len(ranked) == 2
    names = [rc["candidate"]["full_name"] for rc in ranked]
    assert names[0] == "Strong Match"  # higher fit ranked first
    # scores are descending
    assert ranked[0]["score"]["final_score"] >= ranked[1]["score"]["final_score"]
    # reasoning present
    assert ranked[0]["score"]["llm_reasoning"]


def test_scores_persisted_and_stable(client, recruiter_headers):
    job_id = _setup_job_with_two_candidates(client, recruiter_headers)
    first = client.get(
        f"/v1/jobs/{job_id}/ranked-candidates", headers=recruiter_headers
    ).json()
    # second call should return identical (persisted) scores
    second = client.get(
        f"/v1/jobs/{job_id}/ranked-candidates", headers=recruiter_headers
    ).json()
    assert [rc["score"]["final_score"] for rc in first] == [
        rc["score"]["final_score"] for rc in second
    ]


def test_rescore_query_param(client, recruiter_headers):
    job_id = _setup_job_with_two_candidates(client, recruiter_headers)
    forced = client.get(
        f"/v1/jobs/{job_id}/ranked-candidates?rescore=true", headers=recruiter_headers
    )
    assert forced.status_code == 200
    assert len(forced.json()) == 2


def test_candidate_detail(client, recruiter_headers):
    job_id = _setup_job_with_two_candidates(client, recruiter_headers)
    # ensure scores exist
    ranked = client.get(
        f"/v1/jobs/{job_id}/ranked-candidates", headers=recruiter_headers
    ).json()
    cand_id = ranked[0]["candidate"]["id"]
    detail = client.get(f"/v1/candidates/{cand_id}", headers=recruiter_headers)
    assert detail.status_code == 200
    assert detail.json()["candidate"]["id"] == cand_id


# ---- RBAC -----------------------------------------------------------------

def test_recruiter_cannot_delete_candidate(client, recruiter_headers):
    job_id = _setup_job_with_two_candidates(client, recruiter_headers)
    cand_id = client.get(
        f"/v1/jobs/{job_id}/candidates", headers=recruiter_headers
    ).json()[0]["id"]
    r = client.delete(f"/v1/candidates/{cand_id}", headers=recruiter_headers)
    assert r.status_code == 403


def test_admin_can_delete_candidate(client, admin_headers):
    job_id = _setup_job_with_two_candidates(client, admin_headers)
    cand_id = client.get(
        f"/v1/jobs/{job_id}/candidates", headers=admin_headers
    ).json()[0]["id"]
    r = client.delete(f"/v1/candidates/{cand_id}", headers=admin_headers)
    assert r.status_code == 204
    remaining = client.get(f"/v1/jobs/{job_id}/candidates", headers=admin_headers).json()
    assert len(remaining) == 1


# ---- parser: real PDF/DOCX (skips if libs unavailable) --------------------

def test_pdf_and_docx_parsing_roundtrip(tmp_path):
    from app.services.parser import extract_text

    # PDF
    try:
        from reportlab.pdfgen import canvas  # type: ignore

        pdf_path = tmp_path / "r.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(72, 720, "Alice Example")
        c.drawString(72, 700, "Skills: Python, Docker")
        c.save()
        assert "Alice Example" in extract_text(pdf_path)
    except ImportError:
        pytest.skip("reportlab not installed")

    # DOCX
    try:
        import docx  # type: ignore

        docx_path = tmp_path / "r.docx"
        d = docx.Document()
        d.add_paragraph("Bob Example")
        d.add_paragraph("Skills: Java, AWS")
        d.save(str(docx_path))
        assert "Bob Example" in extract_text(docx_path)
    except ImportError:
        pytest.skip("python-docx not installed")
