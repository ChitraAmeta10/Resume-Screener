"""Extract raw text from uploaded resume files (PDF / DOCX / TXT).

Heavy parsing libraries are imported lazily so the module (and the test suite)
loads even if a given parser isn't installed until it's actually needed.
"""
from __future__ import annotations

from pathlib import Path


class UnsupportedFileTypeError(ValueError):
    pass


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix in (".txt", ".md"):
        text = path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {suffix or '(none)'}")
    return _normalize(text)


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber  # lazy

        parts: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception as exc:
        # Fallback to pypdf / pypdf2 if available or basic extraction
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            raise ValueError(f"Failed to extract PDF text: {exc}")


def _extract_docx(path: Path) -> str:
    # 1. Fast, lightweight zero-memory extraction directly from DOCX XML package
    try:
        import xml.etree.ElementTree as ET
        import zipfile

        with zipfile.ZipFile(str(path)) as z:
            xml_content = z.read("word/document.xml")
        tree = ET.fromstring(xml_content)
        texts = []
        for p in tree.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
            p_text = "".join(
                node.text
                for node in p.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                if node.text
            )
            if p_text.strip():
                texts.append(p_text.strip())
        if texts:
            return "\n".join(texts)
    except Exception:
        pass

    # 2. Fallback to python-docx
    try:
        import docx  # lazy

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:
        raise ValueError(f"Failed to extract DOCX text: {exc}")


def _normalize(text: str) -> str:
    lines = [ln.rstrip() for ln in text.replace("\r\n", "\n").split("\n")]
    # collapse runs of blank lines
    out: list[str] = []
    blank = False
    for ln in lines:
        if ln.strip():
            out.append(ln)
            blank = False
        elif not blank:
            out.append("")
            blank = True
    return "\n".join(out).strip()
